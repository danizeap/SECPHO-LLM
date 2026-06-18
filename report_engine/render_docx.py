"""Render a structured Report into a .docx with clean, well-typeset styling.

Until the branded plantilla4.docx is supplied, this uses neutral programmatic
styles isolated here so the template swap is a one-liner later. Sections 1-5 are
populated; Section 6 (Proyectos) appears once the projects data source exists.
"""
from __future__ import annotations

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

BRAND = RGBColor(0x00, 0x86, 0x8C)
DARK = RGBColor(0x1B, 0x1C, 0x20)
MUTED = RGBColor(0x5A, 0x5C, 0x63)


def _ensure_styles(doc) -> None:
    existing = {s.name for s in doc.styles}

    def mk(name, size, bold, color, space_before, space_after):
        st = doc.styles[name] if name in existing else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(space_before)
        st.paragraph_format.space_after = Pt(space_after)

    mk("RG Title", 22, True, BRAND, 0, 14)
    mk("RG H1", 15, True, BRAND, 16, 6)
    mk("RG H2", 12, True, DARK, 10, 4)
    mk("RG Muted", 9.5, False, MUTED, 0, 8)


def _bullet(doc, text: str, level: int = 1) -> None:
    style = "List Bullet" if level == 1 else "List Bullet 2"
    try:
        doc.add_paragraph(text, style=style)
    except KeyError:
        doc.add_paragraph(("   • " if level > 1 else "• ") + text)


def render(report):
    doc = Document()
    _ensure_styles(doc)

    doc.add_paragraph(report.title, style="RG Title")
    doc.add_paragraph("Informe interno de matchmaking y oportunidades · SECPHO", style="RG Muted")

    doc.add_paragraph("Índice", style="RG H1")
    for item in report.indice:
        _bullet(doc, item)

    # 1. Introducción
    doc.add_paragraph("1. Introducción", style="RG H1")
    for para in report.intro:
        doc.add_paragraph(para)

    # 2. Resumen de datos
    doc.add_paragraph(f"2. Resumen de datos de {report.subject_name}", style="RG H1")
    doc.add_paragraph(report.ficha_label, style="RG H2")
    for label, value in report.ficha:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    # 3. Contactos recomendados
    doc.add_paragraph("3. Contactos recomendados", style="RG H1")
    doc.add_paragraph(
        "Contactos del ecosistema SECPHO sugeridos por afinidad de capacidades, "
        "tecnologías clave y áreas de interés. El orden lo determina el modelo de "
        "recomendación: la matemática decide, el informe lo explica."
    )
    if report.contacts:
        for i, c in enumerate(report.contacts, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {c['name']}").bold = True
            extra = " — ".join(x for x in [c.get("socio"), c.get("role")] if x)
            if extra:
                p.add_run(f" · {extra}")
            if c.get("shared_tech"):
                _bullet(doc, f"Tecnologías en común: {c['shared_tech']}", level=2)
            if c.get("shared_sectors"):
                _bullet(doc, f"Sectores en común: {c['shared_sectors']}", level=2)
            if c.get("shared_ambitos"):
                _bullet(doc, f"Ámbitos en común: {c['shared_ambitos']}", level=2)
    else:
        doc.add_paragraph("No se han encontrado contactos recomendados para este perfil en este momento.")

    # 4. Eventos y actividades
    doc.add_paragraph("4. Eventos y actividades", style="RG H1")
    doc.add_paragraph("Próximos eventos recomendados", style="RG H2")
    if report.events_rec:
        for e in report.events_rec:
            p = doc.add_paragraph()
            p.add_run(e["title"]).bold = True
            p.add_run(f"  ({e['score']:.0f}% de afinidad)")
            _bullet(doc, f"Tecnologías: {e['technologies']}", level=2)
            _bullet(doc, f"Sectores: {e['sectors']}", level=2)
            _bullet(doc, f"Ámbitos: {e['ambitos']}", level=2)
            _bullet(doc, f"Ubicación: {e['location']} · Fecha: {e['date']}", level=2)
    else:
        doc.add_paragraph("No hay próximos eventos relevantes que recomendar en este momento.")

    doc.add_paragraph("Histórico de participación en eventos SECPHO", style="RG H2")
    if report.attended:
        for a in report.attended:
            line = f"{a['title']} ({a['date']})"
            if a.get("attendees"):
                line += f" — {a['attendees']}"
            _bullet(doc, line)
    else:
        doc.add_paragraph("Todavía no se registra asistencia a eventos organizados por SECPHO.")

    # 5. Retos tecnológicos
    doc.add_paragraph("5. Retos tecnológicos", style="RG H1")
    doc.add_paragraph("Retos tecnológicos activos recomendados", style="RG H2")
    if report.retos_rec:
        for r in report.retos_rec:
            p = doc.add_paragraph()
            head = f"{r['number']} — {r['title']}" if r.get("number") else r["title"]
            p.add_run(head).bold = True
            if r.get("description"):
                _bullet(doc, r["description"], level=2)
            _bullet(doc, f"Sector(es): {r['sectors']}", level=2)
            _bullet(doc, f"Entidad emisora: {r['issuer']}", level=2)
            _bullet(doc, f"Fecha de cierre: {r['closing']}", level=2)
    else:
        doc.add_paragraph("No se han encontrado retos tecnológicos activos relevantes en este momento.")

    # Emitted/applied retos are entity-level: only shown when there is a socio/company.
    if report.entity_for_retos:
        doc.add_paragraph(f"Retos emitidos por {report.entity_for_retos} a través de SECPHO", style="RG H2")
        if report.retos_emit:
            for r in report.retos_emit:
                head = f"{r['number']} — {r['title']}" if r.get("number") else r["title"]
                _bullet(doc, f"{head} (cierre: {r['closing']})")
        else:
            doc.add_paragraph(f"{report.entity_for_retos} no ha emitido retos tecnológicos a través de SECPHO.")

        doc.add_paragraph(f"Retos en los que {report.entity_for_retos} ha aplicado", style="RG H2")
        if report.retos_appl:
            for r in report.retos_appl:
                head = f"{r['number']} — {r['title']}" if r.get("number") else r["title"]
                _bullet(doc, f"{head} — emisor: {r['issuer']} (cierre: {r['closing']})")
        else:
            doc.add_paragraph(f"{report.entity_for_retos} no consta como aplicante en retos tecnológicos.")

    return doc
