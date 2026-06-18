"""Golden + behaviour tests for report_engine (Phase 1: sections 1-5).

Run: python -m pytest tests/test_report_engine.py
"""
import docx
import pytest

import report_engine as RE
from report_engine import data_access as da
from report_engine import report as rg
from report_engine import scoring


def _valid_target() -> int:
    mem_ids = {int(x) for x in da.members()["member_id"].tolist()}
    for t in da.matches()["target_member_id"].tolist():
        if int(t) in mem_ids:
            return int(t)
    raise RuntimeError("no usable target member in the data")


def _socio_with_members() -> str:
    counts = da.members()["socio"].astype(str).str.lower().value_counts()
    for s in da.socios()["socio"].tolist():
        if counts.get(str(s).lower(), 0) >= 1:
            return s
    raise RuntimeError("no socio with members")


def _paras(path) -> list[str]:
    return [p.text for p in docx.Document(str(path)).paragraphs]


def test_person_report_has_all_sections(tmp_path):
    out = tmp_path / "p.docx"
    RE.generate("person", _valid_target(), str(out))
    assert out.exists() and out.stat().st_size > 0
    text = "\n".join(_paras(out))
    for section in (
        "1. Introducción", "2. Resumen de datos", "3. Contactos recomendados",
        "4. Eventos y actividades", "5. Retos tecnológicos",
    ):
        assert section in text


def test_company_report_generates(tmp_path):
    out = tmp_path / "c.docx"
    RE.generate("company", _socio_with_members(), str(out))
    assert out.exists()
    assert "Ficha de socio" in "\n".join(_paras(out))


def test_output_is_deterministic(tmp_path):
    tid = _valid_target()
    a, b = tmp_path / "a.docx", tmp_path / "b.docx"
    RE.generate("person", tid, str(a))
    RE.generate("person", tid, str(b))
    assert _paras(a) == _paras(b)


def test_no_mojibake_and_has_accents(tmp_path):
    out = tmp_path / "p.docx"
    RE.generate("person", _valid_target(), str(out))
    text = "\n".join(_paras(out))
    assert "�" not in text
    assert any(c in text for c in "áéíóúñ")


def test_contacts_come_only_from_matcher():
    tid = _valid_target()
    report = rg.build_person_report(tid)
    matcher_names = [da._clean(c["candidate_name"]) for c in da.contacts_for_person(tid)]
    assert [c["name"] for c in report.contacts] == matcher_names  # order preserved, none invented


def test_compound_ambito_not_shredded():
    s = da._split_events("New Space, defensa y seguridad, Industria 4.0")
    assert "new space, defensa y seguridad" in s
    assert "industria 4.0" in s
    assert "new space" not in s


def test_attended_dates_are_clean_strings():
    report = rg.build_person_report(_valid_target())
    for a in report.attended:
        assert isinstance(a["date"], str) and a["date"]
        assert "NaT" not in a["date"]


def test_event_scores_are_bounded():
    report = rg.build_person_report(_valid_target())
    for e in report.events_rec:
        assert 0 < e["score"] <= 100


def test_unknown_person_raises():
    with pytest.raises(ValueError):
        rg.build_person_report(999_999_999)


def test_unknown_socio_raises():
    with pytest.raises(ValueError):
        rg.build_company_report("___no_such_socio___")


def test_reto_entity_matching_is_token_exact_not_substring():
    # 'Roca' must NOT match 'ProCareLight' as a substring -> no fabricated reto participation
    assert len(da.retos_applied_by("Roca")) == 0
    assert len(da.retos_issued_by("Roca")) == 0
    # a legitimate applicant still matches
    assert len(da.retos_applied_by("Eurecat")) > 0


def test_attended_dates_are_not_export_stamps():
    # the old bug dated every attended event to the registration-export day (24/25 abril 2025)
    report = rg.build_company_report("AINIA")
    dates = [a["date"] for a in report.attended]
    assert not (dates and all("abril de 2025" in d for d in dates))


def test_vocabulary_canonicalization_matches_across_sources():
    assert da._canon("Robótica & Drones") == da._canon("Robótica y Drones")
    assert da._canon("Sector Farmaceutico") == da._canon("sector farmacéutico")


def test_online_event_requires_topical_overlap():
    empty = {"tech": set(), "sectors": set(), "ambitos": set(), "province": "", "text": ""}
    assert scoring.recommend_events(empty, []) == []  # online bonus alone must not surface an event


def test_generate_bytes_returns_valid_docx():
    import io
    data, fn = RE.generate_bytes("person", _valid_target())
    assert isinstance(data, bytes) and len(data) > 100
    assert fn.startswith("Informe_") and fn.endswith(".docx")
    docx.Document(io.BytesIO(data))  # must open as a valid Word document
    # ASCII-safe filename (no accents/spaces that would break a Content-Disposition header)
    assert fn.encode("ascii")
